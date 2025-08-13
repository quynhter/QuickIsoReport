import os
import re
import random
import pandas as pd
from openpyxl import Workbook
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


def search_excel_file(): # Поиск файла с иходными данными
    print('Начинаю поиск таблицы с исходными данными...')
    files = os.listdir('./')
    for f in files:
        if '.xlsx' in f and f !='Результат.xlsx':
            print(f'Обнаружена подходящая таблица: {f}')
            return f
    return None

FILE_PATH = search_excel_file()


def parse_expression(expr): # Вычисление количества жил
    expr = expr.replace(",", ".").replace("х", "x")
    sections = re.search(r"\d+x\d+(?:\.\d+)?(?:x\d+\.\d+)?", expr)
    cores = re.split(r"x", sections[0])
    total = 1
    
    if len(cores) > 2:
        total = int(cores[0]) * int(cores[1])
    elif len(cores) == 2:
        total = int(cores[0])
    return total



def set_table_borders(table): # Рисовка границ ячеек в таблице в Word
    tbl = table._tbl  # доступ к XML-таблице
    tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')      # тип линии
        border.set(qn('w:sz'), '8')            # толщина (1/8 pt)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')    # цвет
        tblBorders.append(border)

    tblPr.append(tblBorders)


def make_word_result(result_list: list): # Сохранение в Word
    doc = Document()
    last_place = ''

    table = doc.add_table(rows=1, cols=5)
    set_table_borders(table)

    for row in table.rows: # Заполнение и выравнивание ячеек
        for cell in row.cells:
            cell.text = "Текст"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # вертикаль
            for paragraph in cell.paragraphs: # горизонталь
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for result_element in result_list: # Заполнение строк таблицы
        value_title = list(result_element.keys())[0]
        place = result_element[value_title][0]
        values = result_element[value_title][1]

        if last_place == '': # Если самое первое помещение
            row = table.rows[0].cells
            row[0].text = '№ п/п'
            row[1].text = 'Наименование цепи или оборудования'
            row[2].text = 'Марка кабеля, количество жил, сечение (мм²)'
            row[3].text = 'Сопротивление изоляции, МОм'
            row[4].text = 'Заключение о соответствии'
            
            row = table.add_row().cells
            merged_cell = row[0].merge(row[4])
            merged_cell.text = place
            last_place = place
            
            row = table.add_row().cells
            merged_cell = row[0].merge(row[4])
            merged_cell.text = value_title
        
        elif last_place != '' and place != last_place:  # Если помещение сменилось
            row = table.add_row().cells
            merged_cell = row[0].merge(row[4])
            merged_cell.text = place
            last_place = place
            
            row = table.add_row().cells
            merged_cell = row[0].merge(row[4])
            merged_cell.text = value_title
        else: # Если помещение не сменилось
            row = table.add_row().cells
            merged_cell = row[0].merge(row[4])
            merged_cell.text = value_title

        for value in values:
            row = table.add_row().cells
            row[1].text = str(value[0])
            row[2].text = str(value[1])

            i = 1
            for vein in value[2]:
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = f"Жила n×-{i}"
                row[3].text = str(vein)
                row[4].text = "Соответствует"
                i += 1

    doc.save("Результат.docx")


def make_excel_result(result_list: list): # Сохранение в Excel
    wb = Workbook()
    ws = wb.active
    preview_place = ''
    ws.append(['№ п/п', 'Наименование цепи или оборудования', \
              'Марка кабеля, количество жил, сечение (мм²)', \
              'Сопротивление изоляции, МОм', 'Заключение о соответствии'])
    for result_element in result_list:
        value_title = list(result_element.keys())[0]
        place = result_element[value_title][0]
        values = result_element[value_title][1]
        if place == preview_place:
            pass
        else:
            ws.append([place])
            preview_place = place
        ws.append([value_title])
        for value in values:
            ws.append(['', value[0], value[1]])
            i = 1
            for vein in value[2]:
                ws.append([i, f"Жила n×-{i}", '', vein, 'Соответствует'])
                i += 1
    wb.save('Результат.xlsx')



def main():
    try:
        df_dict = pd.read_excel(FILE_PATH, sheet_name=['Трассы', 'Помещения'])
        print('Таблица успешно загружена в программу!')
    except ValueError:
        return print('Таблица не имеет необходимых листов! Пожалуйста, убедитесь, что таблица подготовлена.\nПодсказка: должны быть листы "Трассы" и "Помещения"')
    except FileNotFoundError:
        return print('Программа не смогла найти подходящую таблицу! Убедитесь, что в папке с программой имеется файл с исходными данными.')
    
    df_data_sheet = df_dict['Трассы']
    df_headers_sheet = df_dict['Помещения']

    result_list = []

    search_list = df_headers_sheet['Конец'].to_list()
    for search_element in search_list:
        place = df_headers_sheet[df_headers_sheet['Конец'] == search_element]['Помещение'].to_list()[0]
        result_list.append({search_element: [place, []]})

    col_a = df_data_sheet.columns[0]  # A: Обозначение
    col_c = df_data_sheet.columns[2]  # C: Конец
    col_d = df_data_sheet.columns[3]  # D: Марка и т.п.

    for result_element in result_list:
        value_title = list(result_element.keys())[0]
        for _, row in df_data_sheet.iterrows():
            if str(row[col_c]).strip() == value_title:
                value_a = str(row[col_a]).strip()
                value_d = str(row[col_d]).strip()
                try:
                    count = parse_expression(value_d)
                    rand_numbers = [random.randint(700, 900) for _ in range(count)]

                    result_element[value_title][1].append([value_a, value_d, rand_numbers])
                except Exception as e:
                    print(f"Ошибка при обработке строки: {value_d} — {e}")
    try:
        make_excel_result(result_list=result_list)
        make_word_result(result_list=result_list)
    except PermissionError:
        return print('Программа не смогла открыть файл, т. к. он уже открыт! Пожалуйста, закройте его.')
    print('Программа успешно завершена! Результат записан в файлы: Результат.xlsx и Результат.docx')
    input('Нажмите Enter, чтобы завершить...')


if __name__ == "__main__":
    main()
